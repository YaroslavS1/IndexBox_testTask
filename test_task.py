import sqlite3
import pandas as pd
import numpy as np
import docx


class indexbox():
    def __init__(self):
        # TODO все обернуть в класс и предоставить интерфейсы с помоью которых можно масштабировать и мнять логику
        # TODO в конструктор в качестве параметров принимать год, фактор, путь к бд и прочее
        conn = sqlite3.connect("test.db")  # connect to test.db

        data = pd.read_sql(
            'SELECT * FROM testidprod WHERE partner is NULL AND state is NULL AND bs = 0  AND factor in (1,2)', conn)
        data = data.groupby('id').mean().reset_index()
        data = data.drop(['mir', 'raw', 'bs', 'id', 'country'], axis=1, inplace=False)
        data.loc[data['res'] == "NULL", 'res'] = np.nan
        self.data = data.groupby(['factor', 'year']).sum().reset_index()

        self.year_new = [x for x in range(2006, 2021, 1)]

        for i in self.year_new:
            if i not in self.data[lambda x: x['factor'] == 1]["year"].unique():
                df = pd.DataFrame({'factor': [1], 'year': [int(i)], 'res': [np.nan]})
                self.data = self.data.append(df, sort=True)
            if i not in self.data[lambda x: x['factor'] == 2]["year"].unique():
                df = pd.DataFrame({'factor': [2], 'year': [int(i)], 'res': [np.nan]})
                self.data = self.data.append(df, sort=True)

    def save_xlsx(self):
        """Method to save all DataFrame to format file .xlsx"""

        # ! Добавляем в Dataframe factor 6
        df6 = self.__add_new_factor()

        self.data = pd.concat([self.data, df6], ignore_index=True)
        self.data = self.data.groupby(['factor', 'year']).mean()
        self.data = self.data.rename(columns={'res': 'world'})
        self.data = pd.pivot_table(self.data, values='world', columns=['factor', 'year'], dropna=False)
        self.data.to_excel("report.xlsx")

    def __add_new_factor(self):
        """Method for saving a DataFrame with a factor of 6 to a .docx file"""

        self.data = self.data.sort_values(by=['factor', 'year'])
        col_one_list = self.data['res'].tolist()[:15]
        col_two_list = self.data['res'].tolist()[-15:]
        df6 = pd.DataFrame()
        for i in range(len(self.year_new)):
            df = pd.DataFrame({'factor': [6], 'year': [self.year_new[i]], 'res': [col_one_list[i] / col_two_list[i]]})
            df6 = df6.append(df, sort=True)

        df6 = df6.groupby(['factor', 'year']).mean().reset_index()

        return df6

    @staticmethod
    def cagrf(df):
        """Function to calculate CAGR"""

        i_y = []
        for i in range(len(df['res'])):
            if str(df['res'][i]) != 'nan':
                i_y.append(i)
        return ((df['res'][i_y[-1]] / df['res'][i_y[0]]) ** (1 / (df['year'][i_y[-1]] - df['year'][i_y[0]])) - 1), \
               df['year'][i_y[0]], df['year'][i_y[-1]]

    def save_docx(self):
        """Method for saving a DataFrame with a factor of 6 to a .docx file"""

        df = self.__add_new_factor()
        doc = docx.Document()

        table = doc.add_table(rows=16, cols=3)
        cell = table.cell(1, 0)
        cell.merge(table.cell(15, 0))

        # добавляем заголовок таблицы
        table.cell(0, 0).paragraphs[0].add_run('Factor').bold = True
        table.cell(0, 1).paragraphs[0].add_run('Year').bold = True
        table.cell(0, 2).paragraphs[0].add_run('World Value').bold = True

        cell = table.cell(1, 0)
        cell.text = str(df['factor'][0])

        for i in range(0, 15):
            cell = table.cell(i + 1, 1)
            cell.text = str(df['year'][i])
            cell = table.cell(i + 1, 2)
            cell.text = str(df['res'][i])

        doc.add_paragraph()

        cagr, s_year, e_year = self.cagrf(df)
        if cagr > 0:
            doc.add_paragraph('Factor 6 grew  by avg {:.2%} every year from {} to {}.'.format(cagr, s_year, e_year))
        elif cagr < 0:
            doc.add_paragraph('Factor 6 decreased by avg {:.2%} every year from {} to {}.'.format(cagr, s_year, e_year))

        doc.save('report.docx')


data_indexbox = indexbox()
data_indexbox.save_docx()
data_indexbox.save_xlsx()


